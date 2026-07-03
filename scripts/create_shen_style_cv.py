#!/usr/bin/env python3
"""Create a compact CV_Shen-style DOCX from structured CV content."""

from __future__ import annotations

import argparse
import json
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


FONT = "Times New Roman"
BODY_SIZE = Pt(10)
HEADING_SIZE = Pt(11)
NAME_SIZE = Pt(13)
FONT_DELTA = 0.0
PAGE_WIDTH = 8.5
PAGE_HEIGHT = 13.1889
TOP_MARGIN = 0.7
BOTTOM_MARGIN = 0.7
LEFT_MARGIN = 0.75
RIGHT_MARGIN = 0.75
RIGHT_TAB = Inches(PAGE_WIDTH - LEFT_MARGIN - RIGHT_MARGIN)
SECTION_HEADINGS = {
    "EDUCATION",
    "PROFESSIONAL EXPERIENCE",
    "ACADEMIC EXPERIENCE",
    "SKILLS & AWARDS",
    "TECHNICAL SKILLS",
}


MONTHS = {
    "jan": 1,
    "january": 1,
    "feb": 2,
    "february": 2,
    "mar": 3,
    "march": 3,
    "apr": 4,
    "april": 4,
    "may": 5,
    "jun": 6,
    "june": 6,
    "jul": 7,
    "july": 7,
    "aug": 8,
    "august": 8,
    "sep": 9,
    "sept": 9,
    "september": 9,
    "oct": 10,
    "october": 10,
    "nov": 11,
    "november": 11,
    "dec": 12,
    "december": 12,
}


def sort_key_from_date(value: str) -> tuple[int, int]:
    text = str(value or "").lower()
    if "present" in text or "current" in text:
        return (9999, 12)

    matches: list[tuple[int, int]] = []
    month_pattern = "|".join(sorted(MONTHS, key=len, reverse=True))
    for month, year in re.findall(rf"\b({month_pattern})\.?\s+(\d{{4}})\b", text):
        matches.append((int(year), MONTHS[month]))
    for year, month in re.findall(r"\b(\d{4})[./-](\d{1,2})\b", text):
        matches.append((int(year), int(month)))
    for year in re.findall(r"\b(20\d{2}|19\d{2})\b", text):
        matches.append((int(year), 12))
    return max(matches) if matches else (0, 0)


def sort_reverse_chronological(entries: list[dict]) -> list[dict]:
    return sorted(entries or [], key=lambda item: sort_key_from_date(item.get("date") or item.get("end_date") or ""), reverse=True)


def adjusted_size(size):
    points = size.pt + FONT_DELTA
    if size == BODY_SIZE:
        points = max(points, 9.5)
    return Pt(points)


def set_run(run, *, bold=None, italic=None, size=BODY_SIZE):
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = adjusted_size(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    return run


def set_para_compact(paragraph, *, before=0, after=0):
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = 1


def add_text(paragraph, text, *, bold=False, italic=False, size=BODY_SIZE):
    run = paragraph.add_run(text or "")
    return set_run(run, bold=bold, italic=italic, size=size)


def add_bottom_border(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")


def add_centered_header(doc, personal):
    name = (personal.get("name") or "STUDENT NAME").strip()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_compact(p, after=0)
    add_text(p, name.upper(), bold=True, size=NAME_SIZE)

    contact_items = []
    for key in ("email", "phone", "address"):
        value = personal.get(key)
        if value:
            contact_items.append(str(value).strip())
    for link in personal.get("links", []) or []:
        if link:
            contact_items.append(str(link).strip())

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_compact(p, after=2)
    add_text(p, " | ".join(contact_items), size=BODY_SIZE)


def add_heading(doc, text):
    p = doc.add_paragraph()
    set_para_compact(p, before=5, after=1)
    add_text(p, text.upper(), bold=True, size=HEADING_SIZE)
    add_bottom_border(p)


def add_right_tab_line(doc, left, right="", *, left_bold=True, left_italic=False, right_italic=False, before=0, after=0):
    p = doc.add_paragraph()
    set_para_compact(p, before=before, after=after)
    p.paragraph_format.tab_stops.add_tab_stop(RIGHT_TAB, WD_TAB_ALIGNMENT.RIGHT)
    add_text(p, left or "", bold=left_bold, italic=left_italic)
    if right:
        add_text(p, "\t")
        add_text(p, right, italic=right_italic)
    return p


def add_bullet(doc, text, *, bold_label=False):
    if not text:
        return
    p = doc.add_paragraph(style="List Bullet")
    set_para_compact(p, after=0)
    p.paragraph_format.left_indent = Pt(13)
    p.paragraph_format.first_line_indent = Pt(-7)
    if bold_label and ":" in text:
        label, rest = text.split(":", 1)
        add_text(p, f"{label}:", bold=True)
        add_text(p, rest)
    else:
        add_text(p, text)


def add_education(doc, entries):
    if not entries:
        return
    add_heading(doc, "EDUCATION")
    for index, item in enumerate(sort_reverse_chronological(entries)):
        add_right_tab_line(
            doc,
            item.get("university", ""),
            item.get("location", ""),
            left_bold=True,
            before=1 if index else 0,
        )
        add_right_tab_line(
            doc,
            item.get("degree", ""),
            item.get("date", ""),
            left_bold=False,
            right_italic=True,
        )

        gpa_bits = []
        if item.get("gpa"):
            gpa_bits.append(f"GPA: {item['gpa']}")
        if item.get("ranking"):
            gpa_bits.append(f"Ranking: {item['ranking']}")
        if gpa_bits:
            add_bullet(doc, "; ".join(gpa_bits), bold_label=True)

        coursework = item.get("coursework") or item.get("relevant_coursework")
        if coursework:
            if isinstance(coursework, list):
                coursework = ", ".join(str(x) for x in coursework if x)
            add_bullet(doc, f"Relevant Coursework: {coursework}", bold_label=True)

        honors = item.get("honors") or item.get("awards")
        if honors:
            if isinstance(honors, list):
                honors = ", ".join(str(x) for x in honors if x)
            add_bullet(doc, f"Honors: {honors}", bold_label=True)


def add_experience_section(doc, heading, entries):
    if not entries:
        return
    add_heading(doc, heading)
    for index, item in enumerate(sort_reverse_chronological(entries)):
        title = item.get("organization") or item.get("company") or item.get("project") or item.get("title") or ""
        add_right_tab_line(doc, title, item.get("location", ""), left_bold=True, before=1 if index else 0)
        add_right_tab_line(
            doc,
            item.get("role") or item.get("position") or "",
            item.get("date", ""),
            left_bold=False,
            left_italic=True,
            right_italic=True,
        )
        for bullet in (item.get("bullets") or [])[:5]:
            add_bullet(doc, str(bullet))


def add_skills(doc, skills):
    if not skills:
        return
    has_awards = False
    if isinstance(skills, dict):
        has_awards = any("award" in str(label).lower() and values for label, values in skills.items())
    heading = "SKILLS & AWARDS" if has_awards else "TECHNICAL SKILLS"
    add_heading(doc, heading)
    if isinstance(skills, dict):
        for label, values in skills.items():
            if not values:
                continue
            if isinstance(values, list):
                values = ", ".join(str(x) for x in values if x)
            add_bullet(doc, f"{label}: {values}", bold_label=True)
    elif isinstance(skills, list):
        for item in skills:
            add_bullet(doc, str(item), bold_label=True)


def build_doc(data):
    doc = Document()
    section = doc.sections[0]
    section.start_type = WD_SECTION.NEW_PAGE
    section.page_width = Inches(PAGE_WIDTH)
    section.page_height = Inches(PAGE_HEIGHT)
    section.top_margin = Inches(TOP_MARGIN)
    section.bottom_margin = Inches(BOTTOM_MARGIN)
    section.left_margin = Inches(LEFT_MARGIN)
    section.right_margin = Inches(RIGHT_MARGIN)
    section.header_distance = Inches(0.3)
    section.footer_distance = Inches(0.3)

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = adjusted_size(BODY_SIZE)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = 1

    add_centered_header(doc, data.get("personal", {}))
    add_education(doc, data.get("education", []))
    add_experience_section(doc, "PROFESSIONAL EXPERIENCE", data.get("professional_experience", []))
    add_experience_section(doc, "ACADEMIC EXPERIENCE", data.get("academic_experience", []))
    add_skills(doc, data.get("skills_awards", {}))
    return doc


def safe_filename_part(text: str) -> str:
    text = re.sub(r'[<>:"/\\|?*]+', "", str(text or "").strip())
    text = re.sub(r"\s+", " ", text)
    return text.strip(" .") or "Student"


def resolve_output_path(requested: Path, data: dict) -> Path:
    if requested.suffix.lower() == ".docx":
        return requested
    personal = data.get("personal", {})
    name = personal.get("english_name") or personal.get("name")
    program = data.get("target_program") or data.get("program") or data.get("target")
    base = f"CV-{safe_filename_part(name)}"
    if program:
        base += f"-{safe_filename_part(program)}"
    return requested / f"{base}.docx"


def main() -> int:
    global FONT_DELTA
    parser = argparse.ArgumentParser()
    parser.add_argument("input_json", type=Path)
    parser.add_argument("output_docx", type=Path, help="Output .docx path or output directory.")
    parser.add_argument("--font-delta", type=float, default=0.0, help="Uniform font-size adjustment in points, e.g. 0.5 or -0.5. Body text will not go below 9.5 pt.")
    args = parser.parse_args()

    FONT_DELTA = args.font_delta
    data = json.loads(args.input_json.read_text(encoding="utf-8-sig"))
    args.output_docx = resolve_output_path(args.output_docx, data)
    doc = build_doc(data)
    args.output_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(args.output_docx)
    print(args.output_docx)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
