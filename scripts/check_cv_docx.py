#!/usr/bin/env python3
"""Run lightweight structural and Shen-style checks on a generated CV DOCX."""

from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


EXPECTED_ORDER = [
    "EDUCATION",
    "PROFESSIONAL EXPERIENCE",
    "ACADEMIC EXPERIENCE",
]
SKILLS_HEADINGS = {"SKILLS & AWARDS", "TECHNICAL SKILLS"}
ALL_HEADINGS = set(EXPECTED_ORDER) | SKILLS_HEADINGS


def paragraph_has_bottom_border(paragraph) -> bool:
    p_pr = paragraph._p.pPr
    if p_pr is None:
        return False
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        return False
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        return False
    return bottom.get(qn("w:val")) not in {None, "nil", "none"}


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("docx", type=Path)
    args = parser.parse_args()

    doc = Document(args.docx)
    paragraphs = [p.text for p in doc.paragraphs]
    full_text = "\n".join(paragraphs)
    warnings = []

    if not full_text.strip():
        warnings.append("Document appears to contain no paragraph text.")

    if re.search(r"\{\{[^}]+\}\}|\[[A-Z_ ]{3,}\]|PLACEHOLDER|TODO", full_text):
        warnings.append("Possible placeholder text remains.")

    placeholder_pattern = re.compile(r"\[[^\]]+(if available|provided|placeholder|date|gpa|name|email)[^\]]*\]", re.I)
    if placeholder_pattern.search(full_text):
        warnings.append("Possible bracketed placeholder remains.")

    section_positions = {section: full_text.find(section) for section in EXPECTED_ORDER}
    skills_present = [section for section in SKILLS_HEADINGS if full_text.find(section) >= 0]
    present_sections = [section for section in EXPECTED_ORDER if section_positions[section] >= 0] + skills_present
    if len(present_sections) >= 2:
        positions = [full_text.find(section) for section in present_sections]
        if positions != sorted(positions):
            warnings.append("Section order does not follow the Shen-style default order.")

    heading_paragraphs = [p for p in doc.paragraphs if p.text.strip() in ALL_HEADINGS]
    missing_borders = [p.text.strip() for p in heading_paragraphs if not paragraph_has_bottom_border(p)]
    if missing_borders:
        warnings.append(f"Section heading(s) missing thin bottom border: {missing_borders}.")

    if doc.tables:
        warnings.append(f"Document contains {len(doc.tables)} table(s); verify there are no visible gridlines.")

    empty_bullet_count = sum(1 for p in doc.paragraphs if p.text.strip() in {"-", "•", "·"})
    if empty_bullet_count:
        warnings.append(f"Found {empty_bullet_count} possible empty bullet paragraph(s).")

    bullet_groups = []
    current_count = 0
    in_experience = False
    non_experience_prefixes = (
        "GPA:",
        "Relevant Coursework:",
        "Honors:",
        "Technical Skills:",
        "Research &",
        "Language Skills:",
        "Certifications:",
        "Awards:",
        "Software &",
    )
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text in {"PROFESSIONAL EXPERIENCE", "ACADEMIC EXPERIENCE"}:
            in_experience = True
            current_count = 0
            continue
        if text in SKILLS_HEADINGS:
            if current_count:
                bullet_groups.append(current_count)
            in_experience = False
            current_count = 0
            continue
        if not in_experience:
            continue
        is_bullet = paragraph.style is not None and "Bullet" in paragraph.style.name
        if is_bullet and text and not text.startswith(non_experience_prefixes):
            current_count += 1
        elif current_count and text and not is_bullet:
            bullet_groups.append(current_count)
            current_count = 0
    if current_count:
        bullet_groups.append(current_count)
    overlong_groups = [count for count in bullet_groups if count > 5]
    if overlong_groups:
        warnings.append(f"Found experience bullet group(s) over 5 bullets: {overlong_groups}.")

    print(f"File: {args.docx}")
    print(f"Paragraphs: {len(doc.paragraphs)}")
    print(f"Tables: {len(doc.tables)}")
    print(f"Shen-style sections present: {', '.join(present_sections) if present_sections else 'none'}")
    print(f"Headings with bottom borders: {len(heading_paragraphs) - len(missing_borders)}/{len(heading_paragraphs)}")
    if warnings:
        print("Warnings:")
        for warning in warnings:
            print(f"- {warning}")
        return 1
    print("No lightweight structural warnings found.")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
