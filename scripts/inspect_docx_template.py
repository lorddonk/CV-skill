#!/usr/bin/env python3
"""Inspect a DOCX CV template and write a lightweight JSON report.

This script is intentionally conservative: it extracts structure and style
signals that help an agent recreate or modify a CV template without guessing.
"""

from __future__ import annotations

import argparse
import json
from pathlib import Path

from docx import Document


def paragraph_info(paragraph):
    style = paragraph.style.name if paragraph.style is not None else None
    runs = []
    for run in paragraph.runs:
        if not run.text:
            continue
        runs.append(
            {
                "text": run.text[:120],
                "bold": run.bold,
                "italic": run.italic,
                "underline": run.underline,
                "font": run.font.name,
                "size_pt": run.font.size.pt if run.font.size else None,
            }
        )
    fmt = paragraph.paragraph_format
    return {
        "text": paragraph.text[:240],
        "style": style,
        "alignment": str(paragraph.alignment),
        "left_indent_pt": fmt.left_indent.pt if fmt.left_indent else None,
        "first_line_indent_pt": fmt.first_line_indent.pt if fmt.first_line_indent else None,
        "space_before_pt": fmt.space_before.pt if fmt.space_before else None,
        "space_after_pt": fmt.space_after.pt if fmt.space_after else None,
        "line_spacing": fmt.line_spacing,
        "runs": runs[:8],
    }


def table_info(table):
    rows = []
    for row in table.rows[:8]:
        rows.append([cell.text[:160] for cell in row.cells])
    return {
        "rows": len(table.rows),
        "columns": len(table.columns),
        "style": table.style.name if table.style is not None else None,
        "sample_rows": rows,
    }


def section_info(section):
    return {
        "page_width_in": section.page_width.inches,
        "page_height_in": section.page_height.inches,
        "top_margin_in": section.top_margin.inches,
        "bottom_margin_in": section.bottom_margin.inches,
        "left_margin_in": section.left_margin.inches,
        "right_margin_in": section.right_margin.inches,
        "header_distance_in": section.header_distance.inches,
        "footer_distance_in": section.footer_distance.inches,
    }


def inspect(path: Path) -> dict:
    doc = Document(path)
    return {
        "file": str(path),
        "sections": [section_info(section) for section in doc.sections],
        "paragraph_count": len(doc.paragraphs),
        "paragraphs": [paragraph_info(p) for p in doc.paragraphs[:80]],
        "table_count": len(doc.tables),
        "tables": [table_info(t) for t in doc.tables[:20]],
        "styles": sorted({p.style.name for p in doc.paragraphs if p.style is not None}),
    }


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("template", type=Path)
    parser.add_argument("--out", type=Path)
    args = parser.parse_args()

    report = inspect(args.template)
    payload = json.dumps(report, ensure_ascii=False, indent=2)
    if args.out:
        args.out.write_text(payload, encoding="utf-8")
    else:
        print(payload)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())

